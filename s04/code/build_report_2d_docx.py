# pyright: reportMissingImports=false
"""Rapport Word : 2D keypoints (4 cams) -> Biomecanique. -> batch/REPORT_2D_Biomechanics_EN.docx"""
import os, time as _t
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"C:\Users\21652\Downloads\OpenSimOverView\Vision-Based Optical Simulation"
B = os.path.join(ROOT, "batch"); FIG = os.path.join(B, "report_figs_2d")
doc = Document(); doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)


def H(t, l=1): doc.add_heading(t, level=l)
def P(t, b=False): p = doc.add_paragraph(); r = p.add_run(t); r.bold = b; return p
def BL(t): doc.add_paragraph(t, style="List Bullet")
def img(n, w=6.2, cap=None):
    fp = os.path.join(FIG, n)
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap: c = doc.add_paragraph(); rr = c.add_run(cap); rr.italic = True; rr.font.size = Pt(9); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
def table(csv, idx, nd=3, top=None, cols=None):
    if not os.path.exists(csv): return
    d = pd.read_csv(csv, index_col=0)
    if cols: d = d[[c for c in cols if c in d.columns]]
    if top: d = d.head(top)
    t = doc.add_table(rows=1, cols=len(d.columns) + 1); t.style = "Light Grid Accent 1"; t.rows[0].cells[0].text = idx
    for j, c in enumerate(d.columns): t.rows[0].cells[j + 1].text = str(c)
    for ix, row in d.iterrows():
        cc = t.add_row().cells; cc[0].text = str(ix)
        for j, c in enumerate(d.columns):
            v = row[c]; cc[j + 1].text = ("%.*f" % (nd, v)) if isinstance(v, float) else str(v)
    doc.add_paragraph()


ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("From 2D Keypoints (4 cameras) to Internal Biomechanics"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = s.add_run("Predicting torque, muscle forces, activations and fatigue directly from the 2D image keypoints of "
               "4 cameras - no triangulation, no Vicon, no OpenSim at inference - LOSO, 8 subjects"); rs.italic = True; rs.font.size = Pt(11)
doc.add_paragraph()

H("1. Abstract", 1)
P("This report covers the most direct vision branch: predicting internal biomechanics straight from the 2D image "
  "keypoints (u,v) of RShoulder/RElbow/RWrist in the 4 cameras - WITHOUT triangulation, Vicon, or OpenSim at "
  "inference. Raw 2D pixels alone are weak (mean R2 = 0.63) because depth and perspective are lost. With feature "
  "engineering - per-camera 2D elbow angles and a confidence-weighted MULTI-VIEW FUSION that acts as an implicit "
  "triangulation, plus Butterworth smoothing and cumulative features - the model reaches mean LOSO R2 = 0.85, "
  "close to the explicit-3D model (0.90) and the .mot+.osim pipeline (0.95).")

H("2. Objective", 1)
P("Map the 2D keypoints of the 4 cameras directly to the 13 biomechanical targets, the most deployment-friendly "
  "input (no calibration/triangulation step). Input: for each camera, the 2D (u,v,confidence) of RShoulder, "
  "RElbow, RWrist. Output: torque (1) + activations (4) + forces (4) + fatigue (4).")

H("3. Pipeline and data", 1)
P("Video (4 cams) -> RTMPose 2D keypoints (HALPE-26; RShoulder=6, RElbow=8, RWrist=10) -> features -> LightGBM -> "
  "13 targets. 8 subjects, 12,640 frames; 2D at 50 fps resampled to the label grid. The labels were produced once "
  "by OpenSim (teacher).")

H("4. Inputs and feature engineering", 1)
P("Raw 2D pixels are position/perspective-dependent and lose depth, so they are normalised and, crucially, fused "
  "across views. Features:")
BL("Per camera: shoulder-relative, image-normalised positions; 2D elbow angle a2d = angle(-se2d, we2d); bone "
   "orientations atan2(dv,du); normalised 2D segment lengths; mean keypoint confidence.")
BL("MULTI-VIEW FUSION (the key): a_fused = confidence-weighted mean of the 4 cameras' 2D elbow angles. Averaging "
   "the 4 views approximates the true 3D elbow angle - an IMPLICIT triangulation learned via a feature.")
BL("Temporal: 2 Hz zero-phase Butterworth on a_fused, then qd, qdd; rolling mean/std; cumulative cum_path = "
   "integral |qd| dt (fatigue); time.")
BL("Anthropometry: humerus/forearm mass (subject-aware).")
P("Feature-engineering journey: raw 2D = 0.631 -> with per-camera angles + multi-view fusion + Butterworth = "
  "0.845 -> + Optuna = 0.852. The fusion alone recovers most of the lost depth (activations 0.39 -> 0.86).")
img("fig2d_progression.png", 6.0, "Figure 1 - Progression of the 2D model (dashed: 3D and .mot+.osim references).")

H("4.1 The raw 2D keypoints - where the features start, and why raw pixels are weak", 2)
P("For each of the 4 cameras, RTMPose gives, per frame, the pixel position (u, v) and a confidence for each joint. "
  "Worked example (subject s04, camera 50591643, frame 100):")
BL("RShoulder = (518, 247) px, conf 0.88")
BL("RElbow    = (526, 322) px, conf 0.86")
BL("RWrist    = (539, 380) px, conf 0.76")
P("These are PIXELS in that camera's 900x900 image. Two problems: (a) they depend on where the subject stands and "
  "on the camera pose, so they are not comparable between subjects (they break Leave-One-Subject-Out); (b) a "
  "single camera has no depth - the same 3D pose projects to different pixels depending on the viewing angle "
  "(foreshortening). That is why the raw 2D pixels alone reach only mean R2 = 0.63.")

H("4.2 From 2D pixels to features - step by step (same worked example)", 2)
P("Step 1 - Per-camera 2D bone vectors (subtraction):", True)
BL("se2d = RElbow - RShoulder = (7, 75) px ; we2d = RWrist - RElbow = (13, 58) px.")
P("Step 2 - Normalised 2D lengths (translation/scale-robust):", True)
BL("|se2d|/900 = 0.084 ; |we2d|/900 = 0.066.")
P("Step 3 - Per-camera 2D elbow angle (dot product + arccos):", True)
BL("a2d = arccos( (-se2d . we2d) / (|se2d| |we2d|) ) = 173 deg here -> this camera sees the arm nearly straight.")
P("Step 4 - Bone orientations in the image plane:", True)
BL("se_ori = atan2(75, 7) = 85 deg ; we_ori = atan2(58, 13) = 77 deg (image v points downward).")
P("These four steps are computed for ALL 4 cameras, giving 4 independent 2D views of the same elbow.")

H("4.3 Multi-view fusion = implicit triangulation (the decisive step)", 2)
P("Each camera's 2D elbow angle is a perspective-distorted view of the TRUE 3D angle; one camera alone is "
  "unreliable. Averaging the four cameras, weighted by their confidence, cancels much of the per-view distortion:")
BL("a_fused = sum_c ( conf_c * a2d_c ) / sum_c conf_c.")
P("This is an IMPLICIT triangulation: combining the 4 views recovers a robust angle close to the true 3D elbow "
  "angle, WITHOUT explicit calibration. SHAP shows a_fused is the most important feature, and it is the reason the "
  "2D model jumps from 0.63 to 0.85 - the activations in particular go from 0.39 to 0.86, because they need a "
  "faithful elbow angle. We then apply a 2 Hz zero-phase Butterworth filter to a_fused (clean signal), compute "
  "qd, qdd (velocity/acceleration), the cumulative cum_path = integral |qd| dt (the fatigue driver), and rolling "
  "mean/std; the per-camera angles, lengths and confidences and the subject masses complete the set.")
P("Why each feature (the rule behind it):", True)
BL("a_fused -> the elbow flexion (main DOF) -> drives torque, forces and activations.")
BL("qd, qdd -> the Coriolis and inertial terms of the torque equation.")
BL("cum_path, time -> cumulative history -> the fatigue (an integral of activation).")
BL("per-camera a2d + confidences -> let the model re-weight the views and trust the reliable cameras.")
BL("humerus/forearm mass -> make the model subject-aware (torque scales with body size).")

H("5. Model selection", 1)
table(os.path.join(B, "research_2d.csv"), "Model")

H("6. Hyper-parameter optimisation (Optuna) - top trials", 1)
P("Optuna (TPE + Hyperband) tunes the LightGBM. Top trials:")
table(os.path.join(B, "optuna_2d_trials.csv"), "trial", nd=3, top=8,
      cols=["value", "params_n_estimators", "params_num_leaves", "params_learning_rate", "params_subsample", "params_colsample_bytree", "params_min_child_samples"])

H("7. Explainability (SHAP)", 1)
P("SHAP confirms the engineered signals carry the prediction: the multi-view fused angle (a_fused), the per-camera "
  "angles, the cumulative term and the velocity dominate.")
img("fig2d_xai.png", 5.4, "Figure 2 - SHAP feature importance (2D model).")

H("8. Final model - multi-metric evaluation", 1)
table(os.path.join(B, "metrics_2d_final.csv"), "Target", nd=3)
img("fig2d_metrics.png", 5.8, "Figure 3 - R2 per target (2D final).")
P("Per-subject generalisation (LOSO folds):", True)
table(os.path.join(B, "metrics_2d_per_subject.csv"), "Subject", nd=3)
img("fig2d_per_subject.png", 6.0, "Figure 4 - Per-subject R2 by group.")

H("9. Time-resolved comparison (per muscle)", 1)
img("fig2d_muscles.png", 6.6, "Figure 5 - Per-muscle predicted vs truth over time, held-out subject.")

H("10. The three approaches compared", 1)
P("Across the three vision approaches, accuracy rises as the input gets closer to the biomechanics; feature "
  "engineering (especially multi-view fusion = implicit triangulation) is what makes the 2D approach viable.")
img("fig2d_compare.png", 6.0, "Figure 6 - The three approaches: 2D raw -> 2D+FE -> 3D -> .mot+.osim.")
BL(".mot + .osim (Approach A): mean R2 = 0.952 (cleaned angles + scaled model).")
BL("3D joints (explicit triangulation): 0.904.")
BL("2D keypoints + feature engineering (implicit triangulation): 0.852 (raw 2D was only 0.631).")

H("11. Conclusion", 1)
P("From only the 2D image keypoints of 4 cameras, a LightGBM with multi-view-fusion feature engineering predicts "
  "torque, forces, activations and fatigue of an unseen subject at mean LOSO R2 = 0.85 - with no triangulation, "
  "Vicon or OpenSim at inference. The decisive idea is to fuse the per-camera 2D angles into a single robust angle "
  "(an implicit triangulation), then smooth and accumulate it. This is the lightest deployable pipeline; the "
  "explicit-3D and .mot+.osim pipelines remain more accurate when triangulation/lab data are available.")

H("12. Reproducibility", 1)
BL("Data: batch/<subj>/pose2sim/pose/cam_*_json/*.json (2D keypoints) ; labels in labels_ml.csv.")
BL("Code: model_2d.py (raw 2D), research_2d.py (FE + benchmark + Optuna + XAI), report_2d_build.py.")
BL("Model: batch/model_2d_final/lgbm_2d.joblib. Results: research_2d.csv, metrics_2d_final.csv, optuna_2d_trials.csv, xai_2d.csv.")

out = os.path.join(B, "REPORT_2D_Biomechanics_EN.docx")
for cand in [out, os.path.join(B, "REPORT_2D_Biomechanics_EN_%s.docx" % _t.strftime("%H%M"))]:
    try:
        doc.save(cand); out = cand; break
    except PermissionError:
        continue
print("REPORT:", out, "| headings", len([p for p in doc.paragraphs if p.style.name.startswith("Heading")]), "| images", len(doc.inline_shapes), "| tables", len(doc.tables))
