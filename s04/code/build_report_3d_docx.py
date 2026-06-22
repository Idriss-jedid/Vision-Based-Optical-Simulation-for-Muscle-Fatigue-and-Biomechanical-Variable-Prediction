# pyright: reportMissingImports=false
"""Rapport Word complet : 3D joints -> Biomecanique (vision-only). -> batch/REPORT_3D_Biomechanics_EN.docx"""
import os, time as _t
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"C:\Users\21652\Downloads\OpenSimOverView\Vision-Based Optical Simulation"
B = os.path.join(ROOT, "batch"); FIG = os.path.join(B, "report_figs_3d")
doc = Document(); doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)


def H(t, l=1): doc.add_heading(t, level=l)
def P(t, b=False): p = doc.add_paragraph(); r = p.add_run(t); r.bold = b; return p
def BL(t): doc.add_paragraph(t, style="List Bullet")
def img(n, w=6.2, cap=None):
    fp = os.path.join(FIG, n)
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap: c = doc.add_paragraph(); rr = c.add_run(cap); rr.italic = True; rr.font.size = Pt(9); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
def table(csv, idx, nd=3, top=None, cols=None):
    if not os.path.exists(csv): return
    d = pd.read_csv(csv, index_col=0)
    if cols: d = d[[c for c in cols if c in d.columns]]
    if top: d = d.head(top)
    t = doc.add_table(rows=1, cols=len(d.columns) + 1); t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = idx
    for j, c in enumerate(d.columns): t.rows[0].cells[j + 1].text = str(c)
    for ix, row in d.iterrows():
        cc = t.add_row().cells; cc[0].text = str(ix)
        for j, c in enumerate(d.columns):
            v = row[c]; cc[j + 1].text = ("%.*f" % (nd, v)) if isinstance(v, float) else str(v)
    doc.add_paragraph()


# TITRE
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("From 3D Joint Keypoints to Internal Biomechanics — A Vision-Only Surrogate"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = s.add_run("Predicting torque, muscle forces, activations and fatigue from 3 markerless 3D joints\n(RShoulder, RElbow, RWrist) — no Vicon, no OpenSim at inference — LOSO, 8 subjects"); rs.italic = True; rs.font.size = Pt(11)
doc.add_paragraph(); doc.add_paragraph("Table of contents:");

H("1. Abstract", 1)
P("This report covers the fully vision-only branch: predicting internal biomechanics directly from the 3 markerless "
  "3D joints of the right arm (shoulder, elbow, wrist), produced by Pose2Sim triangulation — WITHOUT Vicon and "
  "WITHOUT OpenSim at inference. After feature engineering (geometry from the 3 joints, 2 Hz Butterworth smoothing, "
  "physics terms, cumulative and rolling features) and Optuna tuning, a LightGBM model reaches a mean LOSO R2 of "
  "~0.90 on unseen subjects. We analyse why torque is the hardest target and how a Butterworth filter raises it "
  "from 0.79 to 0.84, benchmark against time-series transformers (TST/PatchTST/iTransformer/TFT), and report the "
  "honest gap to the .mot + .osim pipeline (Approach A, 0.952).")

H("2. Objective", 1)
P("Replace the lab chain (angles from .mot + scaled model .osim + OpenSim ID/SO/3CC) by a model that maps the raw "
  "3D joints to the 13 biomechanical variables, deployable from video alone.")
BL("Input: 3 joints (RShoulder, RElbow, RWrist), 3D coordinates, from Pose2Sim .trc.")
BL("Output: torque (1) + activations (4) + forces (4) + fatigue (4) = 13, the OpenSim labels.")
BL("No Vicon (scaling from .trc), no OpenSim at inference (the model replaces it).")

H("3. Pipeline (3D -> biomechanics)", 1)
P("Video (4 cams) -> Pose2Sim (RTMPose 2D + triangulation) -> .trc (3D). We take 3 joints; from them we build "
  "geometric features; a LightGBM predicts the 13 biomechanical targets. The labels for training were generated "
  "once by OpenSim (teacher); at inference only the 3 joints are needed.")

H("4. Data", 1)
P("8 subjects (s03..s11), 12,640 frames, dumbbell biceps curls (5 reps, 2 kg). The 3D joints come from the "
  "Butterworth-filtered Pose2Sim .trc (50 fps), resampled to the label time grid. Only RShoulder/RElbow/RWrist are "
  "used (arm26 = right arm, 2 DOF); Hip is used once only to detect the vertical (gravity) axis.")

H("5. Inputs (features) and outputs", 1)
P("From the 3 joints we form the bone vectors SE = elbow-shoulder and WE = wrist-elbow, then derive ~30 features:", True)
BL("Geometry: upper-arm & forearm lengths, shoulder-wrist distance, elbow angle, bone elevations vs vertical, vertical components.")
BL("Kinematics: angular velocity/acceleration (after smoothing), sin/cos of the elbow angle.")
BL("Physics: gravity moment at the elbow, inertial term, their sum.")
BL("Temporal context: rolling mean/std (centred window) and cumulative integrals (path, gravity impulse).")
BL("Anthropometry: humerus/forearm mass (subject-aware).")
P("Outputs: the same 13 OpenSim targets (torque/forces/activations/fatigue).")

H("5.1 The 9 raw coordinates and why they are not used directly", 2)
P("The raw input is the 3D position (in metres, world frame, Y = vertical) of the 3 joints, per frame: "
  "RShoulder (x,y,z), RElbow (x,y,z), RWrist (x,y,z) = 9 numbers per frame, from the Pose2Sim triangulation "
  "(.trc, 50 fps) linearly interpolated to the label time grid (100 Hz). Example (subject s03, frame 0):")
BL("RShoulder = (-0.0649, 1.4185, -0.1193)   (Y=1.42 m = shoulder height)")
BL("RElbow    = (-0.1270, 1.1403, -0.1543)")
BL("RWrist    = (-0.1765, 0.8943, -0.0940)   (lowest -> arm hanging down)")
P("These raw coordinates are NOT fed to the model directly: they depend on where the subject stands and faces "
  "(absolute position/orientation), so they differ between subjects and would break Leave-One-Subject-Out. We "
  "transform them into subject-invariant geometric quantities (lengths, angles), then derivatives, physics, "
  "cumulative and rolling features.")

H("5.2 From 9 coordinates to features — step by step (worked example)", 2)
P("Step 1 - Bone vectors (subtraction):", True)
BL("SE = RElbow - RShoulder = (-0.0621, -0.2782, -0.0350)   (upper-arm vector)")
BL("WE = RWrist - RElbow   = (-0.0495, -0.2460,  0.0603)   (forearm vector)")
P("Step 2 - Lengths (Euclidean norm):", True)
BL("ua_len = |SE| = sqrt(0.0621^2 + 0.2782^2 + 0.0350^2) = 0.287 m (upper arm)")
BL("fa_len = |WE| = 0.258 m (forearm) ; sw_dist = |RWrist - RShoulder| = 0.537 m")
P("Step 3 - Elbow angle (dot product + arccos):", True)
BL("q = arccos( (-SE . WE) / (|SE| |WE|) ) = arccos(-0.936) = 159.4 deg")
BL("At frame 0 the arm is almost straight (interior angle ~159 deg, near 180); q decreases as the elbow flexes. "
   "sin(q)=0.35, cos(q)=-0.94 capture the moment-arm geometry.")
P("Step 4 - Elevations vs vertical (dot product with the up axis):", True)
BL("ua_elev = arccos(SE.up / |SE|) = 166 deg ; fa_elev = arccos(WE.up / |WE|) = 162 deg (both point downward)")
BL("WE_up = WE.up = -0.246 ; wrist_up = (RWrist - RShoulder).up = -0.524 m (hand 0.52 m below the shoulder)")
BL("Why: the gravity moment is proportional to sin(fa_elev): ~0.31 here, rising to 1 when the forearm is horizontal.")
P("Step 5 - Derivatives (finite differences, AFTER Butterworth smoothing):", True)
BL("qd[i] = (q[i+1] - q[i-1]) / (2*dt) ; qdd[i] = (qd[i+1] - qd[i-1]) / (2*dt), with dt = 0.01 s.")
BL("Why: qd enters the Coriolis term, qdd the inertial term M(q)*qdd. qdd is the noisiest -> smoothing first (Step on Butterworth).")
P("Step 6 - Physics features (closed-form):", True)
BL("grav = g*sin(fa_elev)*(m_forearm*fa_len/2 + 2*fa_len)  (gravitational moment at the elbow)")
BL("inertia = (m_forearm*fa_len^2/3 + 2*fa_len^2)*qdd ; phys = grav + inertia")
P("Step 7 - Cumulative features (running sum, for fatigue):", True)
BL("cum_path[i] = sum_{k<=i} |qd[k]|*dt (total angular path) ; cum_grav[i] = sum |grav[k]|*dt (gravitational dose)")
P("Step 8 - Rolling features (moving average, denoising + context):", True)
BL("roll_mean_q[i] = mean(q over a centred 31-frame window) ; roll_std_q[i] = local standard deviation (same for grav, qd).")
P("Plus the anthropometry (humerus/forearm mass) makes the model subject-aware. Total: ~30 features.")

H("5.3 The Butterworth filter explained", 2)
P("A Butterworth filter is a low-pass filter: it lets low frequencies pass and removes high frequencies. The curl "
  "motion is slow (~0.5 Hz); markerless jitter is fast (~3-6 Hz). A 2 Hz cutoff therefore keeps the real motion "
  "and removes the noise.")
BL("Cutoff fc = 2 Hz ; normalised cutoff Wn = fc / (fs/2) = 2/(100/2) = 0.04 (fraction of the Nyquist frequency).")
BL("2nd order -> moderate roll-off; 'Butterworth' = maximally flat passband (no ripple).")
BL("Zero-phase (filtfilt): the filter is applied forward then backward, cancelling the time delay a normal filter "
   "would introduce -> the kinematics are not shifted in time.")
BL("Applied to the angle BEFORE differentiating, because differentiation amplifies high-frequency noise; smoothing "
   "first yields a clean qdd and is the single most effective lever for the torque (0.77 -> 0.83).")

H("6. Feature-engineering journey (chronological, with the math)", 1)
P("This section tells the story in order: from the first weak model to the final one, what we added at each "
  "stage, the math behind each feature, why it helped, and the resulting score. The starting problem was a weak "
  "torque; the journey shows how we fixed it.")

P("Stage 0 - Raw 3D, basic geometry  (mean R2 = 0.842, torque = 0.72).", True)
P("From the 3 joints we form the bone vectors  SE = elbow - shoulder  and  WE = wrist - elbow. Basic features:")
BL("Elbow angle:  q_el = arccos( (-SE . WE) / (|SE| |WE|) ).")
BL("Segment lengths |SE|, |WE| ; bone elevations vs vertical  ua_elev = angle(SE, up), fa_elev = angle(WE, up).")
BL("Angular velocity/acceleration by finite difference:  qd = dq/dt,  qdd = d2q/dt2.")
P("Problem: torque is weak (0.72). Reason: torque = M(q) qdd + C(q,qd) qd + G(q); the inertial term needs qdd "
  "(2nd derivative), the noisiest quantity from markerless 3D.")

P("Stage 1 - Enriched, physics-guided features  (mean R2 = 0.867, torque = 0.79).", True)
BL("Geometry/gravity: sin(q_el), cos(q_el), and a gravitational-load proxy "
   "grav = (m_forearm + 2) * L_forearm * sin(q_sh + q_el)  (weight x lever x sin of forearm-from-vertical).")
BL("Cumulative (for fatigue, which is an integral of activation): cum_path = integral |qd| dt, "
   "cum_grav = integral |grav| dt.")
BL("Temporal context (denoising): rolling mean/std of q_el over a centred window.")
P("Why it helped: the trig terms expose the gravity geometry; the cumulative integrals are the physical drivers "
  "of 3CC fatigue; the rolling mean denoises the angle. Net: +0.025 overall, torque +0.07, fatigue ~0.92.")

P("Stage 2 - Torque deep-dive: diagnosing the bottleneck.", True)
P("The torque equation shows the inertial term M(q) qdd dominates the error because qdd is noisy. We tried "
  "explicit physics features - gravity moment  tau_grav = g sin(alpha) (m_fa L/2 + m_load L)  and inertial term "
  "tau_inertia = (m_fa L^2/3 + m_load L^2) qdd - but SHAP showed they were REDUNDANT (the angle features already "
  "encode gravity). The real issue was the noise in qdd, not missing terms.")

P("Stage 3 - Butterworth 2 Hz smoothing  (THE lever: torque 0.77 -> 0.83, MAE 1.05 -> 0.86 N.m; mean -> 0.895).", True)
P("We apply a zero-phase 2nd-order Butterworth low-pass to the joint angle BEFORE differentiating "
  "(cutoff fc = 2 Hz, normalised Wn = fc / (fs/2), applied with filtfilt for zero phase). It removes the "
  ">2 Hz markerless noise while keeping the curl signal (~0.5 Hz), giving a clean qdd. This is exactly the filter "
  "used for the lab motion. Cleaner kinematics lift not only torque but also activations and forces.")

P("Stage 4 - Optuna tuning  (mean 0.895 -> 0.904, torque -> 0.84).", True)
P("Bayesian tuning (TPE + Hyperband) of the LightGBM on a combined objective (0.5 torque + 0.5 overall) gives the "
  "final model. Overall journey: 0.842 -> 0.867 -> 0.895 -> 0.904 ; torque: 0.72 -> 0.79 -> 0.83 -> 0.84.")
img("fig3d_progression.png", 6.0, "Figure 1 - Progression of the vision-only model: basic 3D -> enriched FE -> +Butterworth -> +Optuna (dashed: Approach A using .mot + .osim).")

H("7. Model selection: GBM vs time-series transformers", 1)
P("We benchmarked gradient boosting against several time-series deep models (TST, PatchTST, iTransformer, "
  "TFT-lite) on the raw 3D windows. On 8 subjects the boosting + feature engineering wins clearly.")
img("fig3d_compare.png", 6.0, "Figure 2 - Model comparison on the 3D->biomechanics task (LOSO).")

H("8. Hyper-parameter optimisation (Optuna)", 1)
P("Optuna (TPE + Hyperband) tunes the LightGBM on a combined objective (0.5 torque + 0.5 overall). Top trials:")
table(os.path.join(B, "optuna_3d_trials.csv"), "trial", nd=3, top=8,
      cols=["value", "params_n_estimators", "params_num_leaves", "params_learning_rate", "params_subsample", "params_colsample_bytree", "params_min_child_samples"])

H("9. Explainability (XAI / SHAP)", 1)
P("SHAP on the 3D features confirms the physics: the denoised elbow angle (rolling mean) drives torque/forces/"
  "activations, and the cumulative features drive fatigue.")
img("fig3d_xai.png", 5.4, "Figure 3 - SHAP feature importance (3D model).")

H("10. Deep-dive: why torque is hardest, and how we fixed it", 1)
P("Torque = M(q)qdd + C qd + G(q). The inertial term needs qdd, the noisiest quantity from markerless 3D. "
  "Diagnosis: torque was limited by noisy qdd and the absence of a Vicon de-bias. Fix: Butterworth 2 Hz smoothing "
  "(clean qdd) + a torque-specific Optuna -> torque R2 0.79 -> 0.84 (MAE 1.12 -> 0.83 N.m). Explicit physics "
  "features (gravity/inertia) were redundant (the angle features already encode gravity, confirmed by SHAP).")
img("fig3d_torque.png", 5.6, "Figure 4 - Torque improvement (vision-only): baseline -> Savitzky-Golay -> Butterworth -> +Optuna.")

H("11. Final model — multi-metric evaluation", 1)
P("LightGBM (multi-output), Butterworth features, Optuna-tuned, LOSO on 8 subjects. Metrics per target:")
table(os.path.join(B, "metrics_3d_final.csv"), "Target", nd=3)
img("fig3d_metrics.png", 5.8, "Figure 5 - R2 per target (final vision-only 3D model).")
P("Per-subject generalisation (the 8 LOSO folds):", True)
table(os.path.join(B, "metrics_3d_per_subject.csv"), "Subject", nd=3)
img("fig3d_per_subject.png", 6.0, "Figure 6 - Per-subject R2 by group (LOSO).")

H("12. Time-resolved comparison (per muscle)", 1)
P("Predicted vs ground truth over time, per muscle (activation/force/fatigue) for a held-out subject.")
img("fig3d_muscles.png", 6.6, "Figure 7 - Per-muscle predicted vs truth, held-out subject.")

H("13. Vision-only (3D joints) vs the .mot + .osim approach — honest gap", 1)
P("Approach A works from the .mot file (joint angles) and the .osim scaled model (the lab pipeline); it reaches "
  "mean R2 = 0.952 (torque 0.937). That .mot/.osim pipeline is the best-case reference. This 3D vision-only model "
  "works directly from the 3 raw joints and reaches ~0.90 (torque ~0.84) - the honest deployment number, since at "
  "inference only the joints are available.")
P("Note on how the .mot/.osim were produced: the angles in .mot were de-biased and the .osim was scaled using the "
  "Vicon ground truth (available in this dataset). That reference cleaning is what makes Approach A the best case; "
  "it is not available in pure vision-only deployment, which is precisely the gap (mostly the angle de-bias). "
  "Important: there is no data leakage - LOSO never sees the test labels; the reference (when used) only cleaned "
  "the inputs, not the targets.")

H("14. Conclusion", 1)
P("From only 3 markerless 3D joints, a LightGBM with physics-guided feature engineering and Butterworth smoothing "
  "predicts torque, forces, activations and fatigue of an unseen subject at mean LOSO R2 ~0.90, with no Vicon and "
  "no OpenSim at inference. Time-series transformers do not beat it at this data scale. Smoothing the kinematics is "
  "the key lever for torque; cumulative features drive fatigue; anthropometry makes the model subject-aware.")

H("15. Reproducibility", 1)
BL("Data: batch/<subj>/pose2sim/pose-3d/*filt_butterworth.trc (3D) ; labels in batch/<subj>/labels_ml.csv.")
BL("Code: build_3d_model.py, research_3d_gbm.py, ts_transformer_3d.py, ts_transformers_3d.py, fix_torque_3d.py, "
   "improve_torque_3d.py, final_3d_model.py, report_3d_build.py.")
BL("Results: metrics_3d_final.csv, metrics_3d_per_subject.csv, optuna_3d_trials.csv, xai_3d.csv, MASTER_comparison.csv.")

out = os.path.join(B, "REPORT_3D_Biomechanics_EN.docx")
for cand in [out, os.path.join(B, "REPORT_3D_Biomechanics_EN_%s.docx" % _t.strftime("%H%M"))]:
    try:
        doc.save(cand); out = cand; break
    except PermissionError:
        continue
print("REPORT:", out, "| headings", len([p for p in doc.paragraphs if p.style.name.startswith("Heading")]), "| images", len(doc.inline_shapes), "| tables", len(doc.tables))
