# pyright: reportMissingImports=false
"""Construit le rapport Word complet -> batch/RAPPORT_ML_Fatigue.docx. biomech env."""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"C:\Users\21652\Downloads\OpenSimOverView\Vision-Based Optical Simulation"
B = os.path.join(ROOT, "batch"); FIG = os.path.join(B, "report_figs")
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)


def H(txt, lvl=1): doc.add_heading(txt, level=lvl)
def P(txt, bold=False, italic=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = bold; r.italic = italic; r.font.size = Pt(size); return p
def bullet(txt): doc.add_paragraph(txt, style="List Bullet")
def img(name, w=6.2, cap=None):
    fp = os.path.join(FIG, name) if not os.path.isabs(name) else name
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap:
            c = doc.add_paragraph(); rr = c.add_run(cap); rr.italic = True; rr.font.size = Pt(9)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def table_from_csv(csv, idx_name="Modèle", ndec=3, hi_col=None):
    d = pd.read_csv(csv, index_col=0)
    t = doc.add_table(rows=1, cols=len(d.columns) + 1); t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells; hdr[0].text = idx_name
    for j, c in enumerate(d.columns): hdr[j + 1].text = str(c)
    for i, (ix, row) in enumerate(d.iterrows()):
        cells = t.add_row().cells; cells[0].text = str(ix)
        for j, c in enumerate(d.columns):
            v = row[c]; cells[j + 1].text = ("%.*f" % (ndec, v)) if isinstance(v, (int, float)) else str(v)
    return t


# ============================ PAGE DE GARDE ============================
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("Prédiction de la biomécanique interne et de la fatigue musculaire\npar apprentissage automatique à partir de la vision markerless")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = s.add_run("Vision-Based Optical Simulation — Approche A (cinématique → torque / forces / activations / fatigue)\n"
               "Modèle ML remplaçant OpenSim à l'inférence · Validation Leave-One-Subject-Out · 8 sujets (Fit3D)")
rs.italic = True; rs.font.size = Pt(12)
doc.add_paragraph()

# ============================ 1. RÉSUMÉ ============================
H("1. Résumé exécutif", 1)
P("Ce travail construit un pipeline markerless complet qui transforme une vidéo de flexion du coude "
  "avec haltère (dumbbell biceps curl) en variables biomécaniques internes — couple articulaire, forces "
  "et activations musculaires, et fatigue — puis entraîne un modèle d'apprentissage automatique capable de "
  "PRÉDIRE ces variables directement depuis la cinématique, REMPLAÇANT ainsi la chaîne OpenSim "
  "(Inverse Dynamics + Static Optimization + modèle de fatigue 3CC) à l'inférence.")
P("Résultats clés (validation Leave-One-Subject-Out, 8 sujets, ~12 640 frames) :", bold=True)
bullet("Meilleur modèle : LightGBM + feature engineering + Optuna → R² moyen = 0.952.")
bullet("Couple : R²=0.94 · Activations : R²=0.96 · Forces : R²=0.96 · Fatigue : R²=0.94.")
bullet("Le feature engineering (features cumulatifs physiques) fait passer la fatigue de R²=0.842 à 0.936.")
bullet("Les modèles séquentiels profonds (LSTM, PatchTST, TST) — même optimisés — NE battent PAS "
       "le gradient boosting sur la fatigue (0.82–0.85 vs 0.94). Analyse théorique fournie (§13).")
bullet("L'explicabilité (XAI/SHAP) confirme la physique : la fatigue est pilotée par les features cumulatifs, "
       "le couple par la charge gravitaire et l'inertie.")

# ============================ 2. OBJECTIF ============================
H("2. Objectif et problématique", 1)
P("Objectif du projet (tel que défini dans l'article support) : prédire le couple articulaire, les forces "
  "musculaires, les activations et la fatigue à partir d'observations non invasives. OpenSim calcule ces "
  "grandeurs par des méthodes biomécaniques (ID, SO, 3CC) coûteuses et nécessitant un modèle musculo-squelettique. "
  "L'idée est d'apprendre la fonction cinématique → biomécanique sur des données générées par OpenSim, "
  "puis de s'en passer à l'inférence (temps réel, déploiement).")
P("Approche A retenue : l'entrée du modèle ML est la CINÉMATIQUE (angles articulaires et leurs dérivées) "
  "augmentée de l'anthropométrie du sujet ; la sortie est l'ensemble des 13 variables biomécaniques. "
  "L'approche B (entrée = points 2D bruts de la vision) a été écartée comme moins fondée physiquement.")

# ============================ 3. PIPELINE ============================
H("3. Pipeline global", 1)
P("La chaîne complète s'enchaîne ainsi :")
bullet("Vidéo multi-caméras (4 cams, 50 fps) → Pose2Sim → points 3D triangulés (.trc).")
bullet("Points 3D → calcul des angles articulaires (épaule, coude) → motion arm26 (.mot).")
bullet("Modèle arm26 mis à l'échelle par sujet (scaling Vicon) → OpenSim ID + SO + 3CC → LABELS.")
bullet("Cinématique + anthropométrie (X) et labels (Y) → dataset ML → entraînement + validation LOSO.")
P("Découverte de calibration majeure (rappel) : Pose2Sim attend translation = −R·T (tvec OpenCV) et non "
  "le centre caméra T ; cette correction a fait passer la corrélation de triangulation de ~0.86 à 0.999.")

# ============================ 4. DONNÉES ============================
H("4. Données (Fit3D, 8 sujets)", 1)
P("Exercice : dumbbell biceps curls, 5 répétitions, ~14–21 s par sujet, haltère 2 kg. 8 sujets "
  "(s03, s04, s05, s07, s08, s09, s10, s11). Vérité terrain Vicon utilisée pour valider la cinématique "
  "(coude r≈0.99, MAE 4–6° ; épaule MAE ~1°).")
P("Modèle : arm26 (2 DDL : élévation épaule, flexion coude ; 7 muscles dont 4 fléchisseurs ; haltère 2 kg soudé). "
  "Mise à l'échelle par sujet (humérus ×0.85 pour s08 le plus petit → ×1.04 pour s10 le plus grand).")

# ============================ 5. LABELS OPENSIM ============================
H("5. Génération des labels OpenSim (théorie)", 1)
P("Inverse Dynamics (ID) — couple articulaire :", bold=True)
P("τ = M(q)·q̈ + C(q,q̇)·q̇ + G(q) − τ_ext. Le couple dépend de la position (gravité G), de la vitesse "
  "(Coriolis C) et de l'accélération (inertie M). Méthode instantanée (sans mémoire).")
P("Static Optimization (SO) — forces et activations :", bold=True)
P("min Σ aₘ²  sous  Σ rₘ(q)·Fₘ = τ (Crowninshield-Brand). Distribue le couple sur les muscles en minimisant "
  "l'effort. Méthode instantanée. Validée pour les mouvements lents (le curl) sur les agonistes.")
P("3CC (Xia & Frey-Law 2008) — fatigue :", bold=True)
P("dMF/dt = F·MA(t) − R·MF(t). La fatigue MF(t) est l'INTÉGRALE de l'activation sur tout l'historique "
  "(avec récupération exponentielle). Grandeur CUMULATIVE — point central pour le choix des features (§10, §13).")

# ============================ 6. DATASET ============================
H("6. Le dataset ML (entrées / sorties)", 1)
P("Fichier : batch/ml_dataset_A.csv — 12 640 frames × 8 sujets. Chaque ligne = 1 frame.")
P("ENTRÉE (X) — 11 features de base :", bold=True)
bullet("Cinématique (7) : q_sh, q_el (angles), qd_sh, qd_el (vitesses), qdd_sh, qdd_el (accélérations), time.")
bullet("Anthropométrie (4, constante/sujet, extraite du modèle scalé) : humerus_mass, forearm_mass, "
       "humerus_len, forearm_len → rend le modèle « subject-aware ».")
P("SORTIE (Y) — 13 labels :", bold=True)
bullet("Couple (1) : elbow_moment (N·m) — ID.")
bullet("Activations (4) : act_{BIClong, BICshort, BRA, BRD_hand} (0–1) — SO.")
bullet("Forces (4) : frc_{…} (N) — SO.")
bullet("Fatigue (4) : MF_{…} (% MF) — 3CC.")

# ============================ 7. MÉTHODO ============================
H("7. Méthodologie & validation", 1)
P("Validation Leave-One-Subject-Out (LOSO) : on entraîne sur 7 sujets, on teste sur le 8e jamais vu, "
  "et on répète pour les 8 sujets. C'est une validation croisée au niveau SUJET — elle mesure la vraie "
  "généralisation et évite toute fuite temporelle (un sujet n'est jamais à la fois en train et en test).")
P("Nature temporelle : couple/forces/activations sont instantanés (per-frame) ; la fatigue est cumulative "
  "(dépend de l'historique) → traitée via des features cumulatifs (§10).")

# ============================ 8. RÉCUP s03 ============================
H("8. Récupération du sujet s03 (débogage approfondi)", 1)
P("Symptôme : la Static Optimization de s03 se bloquait indéfiniment à t=1.32 s (les 7 autres sujets "
  "passaient sans problème).")
P("Démarche de diagnostic :", bold=True)
bullet("La motion de s03 est saine à t=1.32 s (pas de NaN, couple ID ≈ 9.5 N·m, lisse).")
bullet("Longueurs de fibres musculaires normales (pas de singularité géométrique).")
bullet("CAUSE RÉELLE trouvée : le moment arm du BIClong devient NaN à cette configuration précise — "
       "dégénérescence du solver de wrapping (WrapEllipsoid 'BIClonghh' sur la tête humérale). Le NaN "
       "contamine la SO → l'optimiseur boucle.")
P("Correctif : sur tout l'amplitude du curl (0–130°), le path du BIClong NE s'enroule PAS sur cet "
  "ellipsoïde (moment arm identique avec/sans wrap, écart = 0.0000 m). Retirer 'BIClonghh' a donc un "
  "effet biomécanique NUL et supprime la dégénérescence numérique → la SO s'exécute jusqu'au bout "
  "(1413 frames, fails=0). Cohérent avec les 7 autres sujets.")

# ============================ 9. FEATURE ENG ============================
H("9. Feature engineering", 1)
P("11 features dérivées ajoutées (toutes calculées depuis la cinématique + anthropométrie, sans OpenSim) :")
bullet("Géométrie/gravité : sin/cos(q_el), sin/cos(q_sh), grav_load = (m_avant-bras+2)·L·sin(q_sh+q_el).")
bullet("Énergie/vitesse : |q̇_el|, |q̈_el|, q̇_el².")
bullet("Interaction : q_el × m_avant-bras.")
bullet("CUMULATIFS (par sujet, pour la fatigue) : cum_path_el = ∫|q̇_el|dt, cum_grav_imp = ∫|charge|dt.")
P("Note : la feature 'rep' (compteur de répétition) a été RETIRÉE — elle comptait 6 au lieu de 5 "
  "(seuil mal placé) ; les features cumulatifs sont de meilleurs proxys de l'historique de fatigue.")
img("fig_fe_gain.png", 6.3, "Figure 1 — Gain du feature engineering + Optuna (LightGBM, LOSO). La fatigue "
    "passe de 0.842 à 0.936.")

# ============================ 10. BENCH TABULAIRE ============================
H("10. Benchmark des modèles tabulaires", 1)
P("Comparaison en LOSO (Ridge, Random Forest, Extra Trees, MLP, XGBoost, LightGBM) puis tuning Optuna "
  "(optimisation bayésienne TPE + Hyperband ≈ BOHB).")
table_from_csv(os.path.join(B, "bench_tabular.csv"))
img("fig_tabular.png", 6.0, "Figure 2 — R² moyen par modèle tabulaire. Le gradient boosting domine.")
P("Le boosting (LightGBM/XGBoost) domine nettement, notamment sur le couple (0.89 vs 0.75 pour le MLP). "
  "Ridge (linéaire) échoue sur le couple (R²=0.01) → le mapping est fortement non linéaire.")

# ============================ 11. BENCH DEEP ============================
H("11. Benchmark deep learning (per-frame / séquentiel)", 1)
table_from_csv(os.path.join(B, "bench_deep.csv"))
img("fig_deep.png", 6.0, "Figure 3 — Modèles profonds (ANN, 1D-CNN, LSTM, CNN+LSTM, Transformer).")
P("Les modèles profonds (meilleur ≈ 0.87) restent en dessous du boosting (0.91–0.95) sur ce jeu de "
  "données de taille modeste (8 sujets).")

# ============================ 12. PRÉDICTION ============================
H("12. Qualité de prédiction (sujet jamais vu)", 1)
img("fig_pred_timeseries.png", 6.6, "Figure 4 — Prédiction LightGBM vs vérité OpenSim sur un sujet exclu "
    "(LOSO) : couple, activation du biceps, et fatigue, au cours du temps.")

# ============================ 13. TS FATIGUE ============================
H("13. Time-series pour la fatigue : pourquoi LightGBM gagne", 1)
P("La fatigue étant cumulative, on a testé des modèles séquentiels (LSTM, PatchTST, TST) avec fenêtres "
  "glissantes causales, feature engineering étendu et tuning Optuna. Comparaison sur la fatigue :")
table_from_csv(os.path.join(B, "ts_fatigue.csv"), idx_name="Modèle")
P("Modèles séquentiels optimisés (Optuna + fenêtre W tunée) :")
table_from_csv(os.path.join(B, "ts_fatigue_tuned.csv"), idx_name="Modèle")
img("fig_ts_fatigue.png", 6.3, "Figure 5 — Fatigue : approche feature-based (LightGBM) vs séquentielle.")
P("Analyse théorique — pourquoi le feature-based bat le séquentiel :", bold=True)
bullet("Statistiques suffisantes : la 3CC donne MF(t)=∫ activation. Les features cumulatifs calculent "
       "cet intégrale explicitement → on injecte la physique ; le modèle n'a plus rien à « apprendre » de l'historique.")
bullet("Fenêtre trop courte : LSTM/PatchTST voient W≈0.72 s, alors que la fatigue s'accumule sur 14–21 s. "
       "La fenêtre ne contient pas l'information ; la feature cumulative encode tout l'historique depuis t=0.")
bullet("Peu de données : la fatigue est lente → 1 trajectoire par sujet → ~8 trajectoires seulement. "
       "Insuffisant pour des réseaux profonds (milliers de paramètres) ; trivial pour un arbre + feature cumulative.")
bullet("Littérature : les arbres boostés dominent les données tabulaires structurées (Grinsztajn 2022).")
bullet("Biais inductif : les TS profonds capturent des motifs/formes locales ; la fatigue est une « dose totale », "
       "pas un motif local → inadéquation.")
P("Conclusion : « problème temporel » n'impose pas « modèle séquentiel ». Quand la forme de la dépendance "
  "temporelle est connue (ici une intégrale) et les données limitées, encoder l'historique dans une feature "
  "cumulative est plus efficace, plus rapide et plus précis.")

# ============================ 14. XAI ============================
H("14. Explicabilité (XAI)", 1)
P("SHAP (TreeExplainer) — importance moyenne |valeur| par feature et par cible :")
table_from_csv(os.path.join(B, "xai_importance.csv"), idx_name="Feature", ndec=3)
img("fig_xai_heatmap.png", 5.6, "Figure 6 — Carte d'importance SHAP (top features × cibles).")
img("xai_shap_fatigue.png" if os.path.exists(os.path.join(B, "xai_shap_fatigue.png")) else os.path.join(B, "xai_shap_fatigue.png"),
    5.2, "Figure 7 — Features les plus importantes pour la FATIGUE (SHAP).")
P("Lecture (confirme la physique) :", bold=True)
bullet("Couple : dominé par grav_load (0.30), cos(q_el) (0.27), q̈_el (0.24) → gravité + posture + inertie (M·q̈+G).")
bullet("Activations / Forces : dominées par q_el (posture / bras de levier) et q̇_el.")
bullet("Fatigue : dominée par cum_path_el (0.48) et cum_grav_imp (0.21) → les cumulatifs, comme prévu.")
P("Le chemin xai_ts.csv montre que pour le LSTM, l'importance par permutation est très faible et diffuse — "
  "le réseau n'exploite aucune feature de façon nette, ce qui explique sa moindre performance.")

# ============================ 15. RÉSULTATS FINAUX ============================
H("15. Résultats finaux (LOSO, 8 sujets)", 1)
P("Meilleur modèle global : LightGBM + features engineered + Optuna.", bold=True)
adv = pd.read_csv(os.path.join(B, "ml_advanced.csv"), index_col=0)
t = doc.add_table(rows=1, cols=6); t.style = "Light Grid Accent 1"
for j, h in enumerate(["Configuration", "R² moyen", "Couple", "Activations", "Forces", "Fatigue"]):
    t.rows[0].cells[j].text = h
for ix in ["LGBM_base", "LGBM_eng", "LGBM_eng_optuna"]:
    row = adv.loc[ix]; c = t.add_row().cells
    name = {"LGBM_base": "Base (11 feat)", "LGBM_eng": "+ Feature eng. (22)", "LGBM_eng_optuna": "+ Optuna"}[ix]
    c[0].text = name
    for j, col in enumerate(["mean", "torque", "activations", "forces", "fatigue"]):
        c[j + 1].text = "%.3f" % row[col]

# ============================ 16. CHRONOLOGIE ============================
H("16. Chronologie : ce qui a été ajouté / amélioré / retiré", 1)
P("Ajouté :", bold=True)
bullet("Extraction des labels OpenSim (ID+SO+3CC) pour 8 sujets → dataset ML.")
bullet("Features anthropométriques (masses + longueurs) → modèle subject-aware.")
bullet("Feature engineering (sin/cos, grav_load, cumulatifs) → +0.03 global, +0.09 fatigue.")
bullet("Tuning Optuna (BOHB-like) sur LightGBM et sur les modèles séquentiels.")
bullet("Benchmark large : tabulaires (6) + deep (5) + séquentiels fatigue (LSTM/PatchTST/TST).")
bullet("XAI : SHAP + importance par permutation.")
bullet("Récupération de s03 via correction du wrap BIClonghh.")
P("Retiré / corrigé :", bold=True)
bullet("Feature 'rep' (compteur erroné 6≠5) — retirée.")
bullet("Wrap 'BIClonghh' pour s03 (effet nul, supprime le blocage SO).")
bullet("Approche B (vision 2D brute) — écartée au profit de l'approche A.")

# ============================ 17. CONCLUSION ============================
H("17. Conclusion & perspectives", 1)
P("Un modèle LightGBM, alimenté par la cinématique markerless + l'anthropométrie + des features physiques "
  "cumulatives, prédit le couple, les forces, les activations et la fatigue d'un sujet jamais vu avec un "
  "R² moyen de 0.952 (LOSO), remplaçant la chaîne OpenSim ID+SO+3CC à l'inférence. Le feature engineering "
  "guidé par la physique (intégrales de charge) est la clé de la prédiction de fatigue ; les modèles "
  "séquentiels profonds n'apportent pas de gain sur ce volume de données.")
P("Perspectives :", bold=True)
bullet("Augmenter le nombre de sujets (les modèles profonds pourraient alors devenir compétitifs).")
bullet("Validation des activations contre l'EMG (références littérature) pour certifier les labels.")
bullet("Déploiement temps réel (les features cumulatives se calculent en ligne).")

# ============================ 18. ANNEXES ============================
H("18. Annexes — fichiers et scripts", 1)
bullet("Dataset : batch/ml_dataset_A.csv (12 640 × 26).")
bullet("Résultats : bench_tabular.csv, bench_deep.csv, ml_advanced.csv, ts_fatigue.csv, ts_fatigue_tuned.csv.")
bullet("XAI : xai_importance.csv, xai_ts.csv.")
bullet("Scripts : extract_labels_all.py, add_morpho.py, recover_s03.py, ml_advanced.py, "
       "bench_tabular.py, bench_deep.py, ts_fatigue.py, ts_fatigue_tuned.py.")

out = os.path.join(B, "RAPPORT_ML_Fatigue.docx"); doc.save(out)
print("RAPPORT écrit :", out)
print("sections:", len([p for p in doc.paragraphs if p.style.name.startswith("Heading")]),
      "| images:", len(doc.inline_shapes))
