from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add an external hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_pr.append(run_style)
    run.append(run_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def add_picture_if_exists(doc: Document, image_path: Path, width: Inches) -> None:
    if image_path.exists():
        doc.add_picture(str(image_path), width=width)
        doc.paragraphs[-1].alignment = 1  # centered
    else:
        doc.add_paragraph(f"[Missing image: {image_path}]")


def main() -> None:
    docs_dir = Path(__file__).resolve().parents[1] / "Docs"
    txt_path = docs_dir / "what i doing now .txt"
    out_docx = docs_dir / "what_i_doing_now.docx"

    figure_postures = docs_dir / "figures" / "representative_postures.png"
    mermaid_dir = docs_dir / "figures" / "mermaid"
    fig_mermaid_extended = mermaid_dir / "extended_model.png"
    fig_mermaid_loaded = mermaid_dir / "loaded_model.png"
    fig_mermaid_workflow = mermaid_dir / "opensim_workflow.png"
    fig_mermaid_data_role = mermaid_dir / "data_role.png"
    video_path = docs_dir / "figures" / "Opensim__VD.mp4"

    if not txt_path.exists():
        raise FileNotFoundError(f"Missing source txt file: {txt_path}")

    lines = txt_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    doc.add_heading("OpenSim Upper-Limb Modeling and Data Pipeline", level=0)
    doc.add_paragraph("Converted from: what i doing now .txt")

    inserted_posture_figure = False
    inserted_video_link = False

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        # Numbered headings like "2.1 ...", "3.5 ...", etc.
        m = re.match(r"^(\d+(?:\.\d+)*)\s+(.+)$", stripped)
        if m:
            section_no = m.group(1)
            title = m.group(2)
            level = min(section_no.count(".") + 1, 4)
            doc.add_heading(f"{section_no} {title}", level=level)

            if stripped == "3.5 OpenSim Simulation Workflow":
                doc.add_paragraph("Rendered Mermaid diagram:")
                add_picture_if_exists(doc, fig_mermaid_workflow, Inches(6.7))

            if stripped == "4. Role of the Generated Data in the Proposed Framework":
                doc.add_paragraph("Rendered Mermaid diagram:")
                add_picture_if_exists(doc, fig_mermaid_data_role, Inches(6.5))

            continue

        # Keep all content text as-is so wording is preserved.
        doc.add_paragraph(line)

        # Inject rendered Mermaid images at corresponding markers.
        if stripped == "Mermaid Diagram — Extended Model Construction":
            doc.add_paragraph("Rendered Mermaid diagram:")
            add_picture_if_exists(doc, fig_mermaid_extended, Inches(6.5))

        if stripped == "Mermaid Diagram — Loaded Model Construction":
            doc.add_paragraph("Rendered Mermaid diagram:")
            add_picture_if_exists(doc, fig_mermaid_loaded, Inches(6.0))

        # Add provided posture figure near Figure 2 mention.
        if ("illustrated in Figure 2" in stripped) and (not inserted_posture_figure):
            doc.add_paragraph("Figure 2 image:")
            add_picture_if_exists(doc, figure_postures, Inches(6.5))
            inserted_posture_figure = True

        # Add video hyperlink after video caption sentence.
        if stripped.startswith("Video 1. OpenSim visualization") and (not inserted_video_link):
            p = doc.add_paragraph("Video file: ")
            if video_path.exists():
                add_hyperlink(p, "Opensim__VD.mp4", video_path.resolve().as_uri())
                doc.add_paragraph(f"Local path: {video_path}")
            else:
                doc.add_paragraph(f"[Missing video: {video_path}]")
            inserted_video_link = True

    # If caption trigger wasn't encountered for any reason, still add assets once.
    if not inserted_posture_figure:
        doc.add_heading("Figure Asset", level=3)
        add_picture_if_exists(doc, figure_postures, Inches(6.5))

    if not inserted_video_link:
        doc.add_heading("Video Asset", level=3)
        p = doc.add_paragraph("Video file: ")
        if video_path.exists():
            add_hyperlink(p, "Opensim__VD.mp4", video_path.resolve().as_uri())
            doc.add_paragraph(f"Local path: {video_path}")
        else:
            doc.add_paragraph(f"[Missing video: {video_path}]")

    doc.save(out_docx)
    print(f"Wrote: {out_docx}")


if __name__ == "__main__":
    main()
